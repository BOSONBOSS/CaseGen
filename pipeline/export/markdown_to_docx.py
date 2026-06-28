"""Convert Markdown to Word (.docx) with IFQM styling."""

import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def _set_document_styles(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _add_table_from_markdown(doc: Document, lines: list):
    if len(lines) < 2:
        return
    header = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        if line.strip().startswith("|"):
            rows.append([c.strip() for c in line.strip("|").split("|")])

    if not header:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for i, h in enumerate(header):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row[: len(header)]):
            table.rows[r_idx + 1].cells[c_idx].text = cell


def markdown_to_docx(markdown: str) -> bytes:
    """Convert Markdown string to .docx bytes."""
    doc = Document()
    _set_document_styles(doc)

    lines = markdown.splitlines()
    i = 0
    table_buffer = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            continue
        elif table_buffer:
            _add_table_from_markdown(doc, table_buffer)
            table_buffer = []

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            if p.runs:
                p.runs[0].font.name = "Times New Roman"
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            if p.runs:
                p.runs[0].font.name = "Times New Roman"
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "":
            pass
        else:
            text = line.strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        i += 1

    if table_buffer:
        _add_table_from_markdown(doc, table_buffer)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
